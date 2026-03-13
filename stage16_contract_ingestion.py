#!/usr/bin/env python3
"""
Stage 16 — Contract Ingestion
==============================
Converts raw contract files (PDF, DOCX, TXT) into the stage4_clauses.json
format consumed by the rest of the analysis pipeline.

Processing pipeline
-------------------
  1. Extract raw text blocks (with page numbers where available)
  2. Normalise whitespace
  3. Classify blocks: heading / body / bullet / table / toc / skip
  4. Segment clauses from block stream
  5. Assign CL-NNN IDs and infer layout_type

Output schema
-------------
  [
    {
      "clause_id"  : "CL-001",
      "page"       : 3,
      "layout_type": "paragraph",   # paragraph | bullet_list | table | heading
      "text"       : "..."
    }
  ]

Usage
-----
  python stage16_contract_ingestion.py --contract contract.pdf
  python stage16_contract_ingestion.py --contract contract.docx --output clauses.json
  python stage16_contract_ingestion.py --contract contract.txt  --quiet
"""

from __future__ import annotations

import argparse
import io
import json
import re
import sys
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator

import regex                         # pip install regex

# ── Lazy library imports (guarded with clear error messages) ──────────────────

def _import_pdfminer():
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import (
            LAParams, LTTextBox, LTTextBoxHorizontal,
            LTTextLine, LTChar, LTFigure, LTPage,
        )
        return extract_pages, LAParams, LTTextBox, LTTextLine, LTChar, LTFigure, LTPage
    except ImportError as e:
        print(
            f"[ERROR] pdfminer.six is required for PDF ingestion.\n"
            f"        Install with: pip install pdfminer.six\n"
            f"        Detail: {e}",
            file=sys.stderr,
        )
        sys.exit(2)


def _import_docx():
    try:
        import docx as _docx
        return _docx
    except ImportError as e:
        print(
            f"[ERROR] python-docx is required for DOCX ingestion.\n"
            f"        Install with: pip install python-docx\n"
            f"        Detail: {e}",
            file=sys.stderr,
        )
        sys.exit(2)


# ── Constants ─────────────────────────────────────────────────────────────────

MIN_CLAUSE_CHARS = 20        # discard shorter candidates
MAX_HEADING_CHARS = 200      # headings longer than this are treated as body
TOC_LINE_RATIO    = 0.50     # fraction of dot-leader/page-ref lines → TOC block
CAPS_WORD_RATIO   = 0.70     # fraction of uppercase words → ALL-CAPS heading

# Regex patterns
_RE_NUMBERED_HDG = regex.compile(
    r"""
    ^\s*                              # optional leading space
    (?:
        §\s*\d{1,3}[a-z]?            # § 1  § 25b  § 1a  (German paragraphs)
        |
        Art(?:icle|\.?)\s*\d{1,3}    # Article 1  Art. 5  Art 12  (intl contracts)
        |
        (?:\d{1,2}\.){1,4}\d{0,2}    # 1.  /  1.1  /  1.2.3  /  1.2.3.4
        |
        [IVXLCDM]{1,6}\.             # roman numerals  I.  II.  XIV.
    )
    [\s\-–]{0,4}                     # separator (space, dash, en-dash)
    [A-ZÜÄÖÑA-Z\"]                   # starts with uppercase (incl. German Umlauts)
    """,
    regex.VERBOSE | regex.MULTILINE,
)

_RE_TOC_LINE = regex.compile(
    r"""
    .{3,}                  # some text
    (?:
        \.{3,}\s*\d{1,3}   # dot leaders:  Text ......  12
        |
        \s{3,}\d{1,3}      # wide spaces:  Text       12
        |
        \t\d{1,3}          # tab + page number
    )
    \s*$
    """,
    regex.VERBOSE,
)

_RE_PAGE_ONLY    = regex.compile(r'^\s*[-–—]?\s*\d{1,3}\s*[-–—]?\s*$')
_RE_BULLET_LINE  = regex.compile(r'^\s*(?:[•·▪▸►‣◦\-–—\*]|\d+[.)]\s)\s+\S')
_RE_PIPE_ROW     = regex.compile(r'\|')
_RE_COLON_HDG    = regex.compile(r'^[^:\n\r]{5,100}:\s*$')

# Words that signal document metadata lines to skip
_SKIP_PATTERNS = regex.compile(
    r"""
    ^(?:
        (?:confidential|draft|version|revision|prepared\s+by|date:|document\s+id|
           page\s+\d|copyright|\u00a9|\bproprietary\b)
        [\s:—\-]*
    )
    """,
    regex.VERBOSE | regex.IGNORECASE,
)

# TOC section titles (any line that IS just these)
_TOC_TITLES = regex.compile(
    r'^(?:table\s+of\s+contents?|contents?|index|inhaltsverzeichnis)\s*$',
    regex.IGNORECASE,
)


# ── Data model ────────────────────────────────────────────────────────────────

@dataclass
class RawBlock:
    """A contiguous piece of text from the source document."""
    text:     str
    page:     int                # 1-indexed; 1 if unknown
    kind:     str = "body"       # body | heading | table | toc | skip
    x0:       float = 0.0        # left edge  (PDF only)
    y0:       float = 0.0        # bottom edge (PDF only)
    font_size: float = 0.0       # dominant font size (PDF only)
    is_bold:  bool = False       # dominant bold flag (PDF only)


@dataclass
class ClauseCandidate:
    blocks:     list[RawBlock] = field(default_factory=list)
    page:       int = 1

    @property
    def text(self) -> str:
        return _merge_blocks(self.blocks)

    @property
    def is_empty(self) -> bool:
        return len(self.text.strip()) < MIN_CLAUSE_CHARS


# ── Text normalisation ────────────────────────────────────────────────────────

def _normalise(text: str) -> str:
    """Normalise Unicode, fix common OCR artifacts, collapse whitespace."""
    # Unicode normalisation (NFC handles ligatures, composed chars)
    text = unicodedata.normalize("NFC", text)
    # Replace smart quotes and dashes with ASCII equivalents
    for src, dst in [
        ("\u2018", "'"), ("\u2019", "'"),
        ("\u201c", '"'), ("\u201d", '"'),
        ("\u2013", "-"), ("\u2014", "-"),
        ("\u00a0", " "),   # non-breaking space
        ("\ufffd", ""),    # replacement char (bad OCR)
    ]:
        text = text.replace(src, dst)
    # Collapse multiple spaces on a single line (keep newlines)
    lines = []
    for line in text.splitlines():
        line = re.sub(r'  +', ' ', line).rstrip()
        lines.append(line)
    # Collapse 3+ blank lines to 2
    text = "\n".join(lines)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _merge_blocks(blocks: list[RawBlock]) -> str:
    """Merge block texts with appropriate separators."""
    if not blocks:
        return ""
    parts: list[str] = []
    for b in blocks:
        t = b.text.strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


# ── Block classification ──────────────────────────────────────────────────────

def _is_toc_block(text: str) -> bool:
    """True if this block looks like a table-of-contents section."""
    if _TOC_TITLES.match(text.strip()):
        return True
    lines = [l for l in text.splitlines() if l.strip()]
    if not lines:
        return False
    toc_lines = sum(1 for l in lines if _RE_TOC_LINE.match(l.strip()))
    return (toc_lines / len(lines)) >= TOC_LINE_RATIO


def _is_all_caps_heading(text: str) -> bool:
    """True if the text is a short ALL-CAPS heading line."""
    stripped = text.strip()
    if "\n" in stripped or len(stripped) > MAX_HEADING_CHARS:
        return False
    words = [w for w in regex.sub(r'[^\w\s]', '', stripped).split()
             if len(w) > 2]
    if len(words) < 1 or len(words) > 12:
        return False
    upper = sum(1 for w in words if w.isupper())
    return (upper / len(words)) >= CAPS_WORD_RATIO


def _is_numbered_heading(text: str) -> bool:
    """True if the block starts with a numbered-section heading."""
    return bool(_RE_NUMBERED_HDG.match(text.strip()))


def _is_colon_heading(text: str) -> bool:
    """True if the block is a single colon-terminated header line."""
    stripped = text.strip()
    return bool(_RE_COLON_HDG.match(stripped)) and "\n" not in stripped


def _is_page_artefact(text: str) -> bool:
    """True if the text is a bare page number or document-metadata noise."""
    stripped = text.strip()
    if _RE_PAGE_ONLY.match(stripped):
        return True
    if _SKIP_PATTERNS.match(stripped):
        return True
    if len(stripped) < 4:
        return True
    return False


def _classify_block(block: RawBlock) -> str:
    """Assign a block kind in place; return the kind string."""
    t = block.text.strip()
    if not t or _is_page_artefact(t):
        block.kind = "skip"
    elif _is_toc_block(t):
        block.kind = "toc"
    elif _is_numbered_heading(t) and len(t) <= MAX_HEADING_CHARS:
        block.kind = "heading"
    elif _is_all_caps_heading(t):
        block.kind = "heading"
    elif _is_colon_heading(t):
        block.kind = "heading_colon"
    else:
        block.kind = "body"
    return block.kind


# ── Layout type inference ─────────────────────────────────────────────────────

def _infer_layout(text: str) -> str:
    """
    Classify the visual layout of a clause's merged text.

    Returns one of: paragraph | bullet_list | table | heading
    """
    lines = [l for l in text.splitlines() if l.strip()]
    if not lines:
        return "paragraph"

    # Table: significant fraction of lines contain pipe characters
    pipe_lines = sum(1 for l in lines if _RE_PIPE_ROW.search(l))
    if pipe_lines >= 2 and pipe_lines / len(lines) >= 0.35:
        return "table"

    # Also detect tab/space-aligned multi-column tables
    if len(lines) >= 3:
        tab_lines = sum(1 for l in lines if '\t' in l or re.search(r'  {3,}', l))
        if tab_lines / len(lines) >= 0.6:
            # Check for consistent column count
            split_counts = [len(re.split(r'\t|  {3,}', l)) for l in lines[:6]]
            if len(set(split_counts)) <= 2 and split_counts[0] >= 2:
                return "table"

    # Bullet list: significant fraction of lines are bullet items
    bullet_lines = sum(1 for l in lines if _RE_BULLET_LINE.match(l))
    if bullet_lines >= 2 and bullet_lines / len(lines) >= 0.30:
        return "bullet_list"
    # Also: single colon-header line followed by bullets
    if len(lines) >= 2 and _is_colon_heading(lines[0]):
        remaining = lines[1:]
        if remaining and sum(1 for l in remaining if _RE_BULLET_LINE.match(l)) / len(remaining) >= 0.5:
            return "bullet_list"

    return "paragraph"


# ═══════════════════════════════════════════════════════════════════════════════
# PDF EXTRACTOR
# ═══════════════════════════════════════════════════════════════════════════════

def extract_pdf(path: Path) -> list[RawBlock]:
    """
    Extract text blocks from a PDF with page numbers.

    Uses pdfminer.six:
    - Iterates LTPage objects (1-indexed page numbers)
    - Collects LTTextBox instances as individual blocks
    - Captures dominant font size from LTChar for heading detection
    - Identifies likely header/footer zones (top/bottom 8% of page)
    """
    extract_pages, LAParams, LTTextBox, LTTextLine, LTChar, LTFigure, LTPage = \
        _import_pdfminer()

    laparams = LAParams(
        line_overlap   = 0.5,
        char_margin    = 2.0,
        line_margin    = 0.6,   # slightly looser than default to keep para together
        word_margin    = 0.1,
        boxes_flow     = 0.5,
        detect_vertical= False,
    )

    blocks: list[RawBlock] = []

    for page_num, page_layout in enumerate(
        extract_pages(str(path), laparams=laparams), start=1
    ):
        page_height = page_layout.height if hasattr(page_layout, 'height') else 841
        margin_v    = page_height * 0.08     # top/bottom 8% = header/footer zone

        # Collect all text boxes on this page with position info
        for element in page_layout:
            if not isinstance(element, LTTextBox):
                continue

            raw_text = element.get_text()
            if not raw_text.strip():
                continue

            # Get dominant font size from LTChar elements
            font_sizes: list[float] = []
            bold_chars:  int = 0
            total_chars: int = 0
            for line in element:
                if isinstance(line, LTTextLine):
                    for char in line:
                        if isinstance(char, LTChar):
                            font_sizes.append(char.size)
                            total_chars += 1
                            fn = (char.fontname or "").lower()
                            if any(b in fn for b in ("bold", "black", "heavy", "-b-", "bd")):
                                bold_chars += 1

            dom_size  = (sum(font_sizes) / len(font_sizes)) if font_sizes else 0.0
            is_bold   = (bold_chars / total_chars > 0.5) if total_chars else False

            # Identify header/footer zone by y-position
            y0 = element.y0
            y1 = element.y1
            in_margin = (y1 < margin_v) or (y0 > page_height - margin_v)

            block = RawBlock(
                text      = raw_text,
                page      = page_num,
                x0        = element.x0,
                y0        = y0,
                font_size = dom_size,
                is_bold   = is_bold,
            )

            # Mark structural kind based on position + content
            if in_margin:
                block.kind = "skip"
            else:
                _classify_block(block)

            # Promote based on PDF font cues (large or bold = likely heading)
            if block.kind == "body" and dom_size > 0:
                # Compare against a baseline we'll compute below
                block.kind = "body"     # refined in post-pass below

            blocks.append(block)

    # Post-pass: compute per-page median font size and promote large/bold short
    # blocks to headings
    if blocks:
        all_sizes = [b.font_size for b in blocks if b.font_size > 0]
        if all_sizes:
            all_sizes.sort()
            median_size = all_sizes[len(all_sizes) // 2]
            for b in blocks:
                if b.kind == "body" and b.font_size > median_size * 1.15:
                    text_lines = b.text.strip().splitlines()
                    # Only a short block can be a heading based on font alone
                    if len(text_lines) <= 3 and len(b.text.strip()) <= MAX_HEADING_CHARS:
                        _classify_block(b)   # re-classify with full rules
                        if b.kind == "body" and (b.font_size > median_size * 1.2 or b.is_bold):
                            b.kind = "heading"

    return blocks


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX EXTRACTOR
# ═══════════════════════════════════════════════════════════════════════════════

def extract_docx(path: Path) -> list[RawBlock]:
    """
    Extract text blocks from a DOCX with estimated page numbers.

    - Headings detected from paragraph style (Heading 1-9, TOC Heading)
    - Tables extracted as a single pipe-separated block per table
    - Page breaks tracked via paragraph.contains_page_break
    - TOC paragraphs identified from TOC 1-9 styles
    """
    _docx = _import_docx()
    doc   = _docx.Document(str(path))
    blocks: list[RawBlock] = []
    current_page = 1

    # python-docx iterates body elements in document order, but
    # doc.paragraphs and doc.tables are separate — use iter_block_items
    # to get them interleaved.
    from docx.oxml.ns import qn

    def _iter_body_items(document):
        """Yield (tag, element) for body paragraphs and tables in order."""
        body = document.element.body
        for child in body.iterchildren():
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                yield "para", _docx.text.paragraph.Paragraph(child, document)
            elif tag == "tbl":
                yield "table", _docx.table.Table(child, document)

    for item_type, item in _iter_body_items(doc):

        if item_type == "para":
            para = item
            text = para.text.strip()

            # Track page breaks inside paragraphs
            if para.contains_page_break:
                current_page += 1

            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # Determine kind from style
            if regex.match(r'(?i)^heading\s*[1-9]', style_name):
                kind = "heading"
            elif regex.match(r'(?i)^toc\s', style_name):
                kind = "toc"
            elif regex.match(r'(?i)^(header|footer)', style_name):
                kind = "skip"
            else:
                kind = None   # determine from content

            block = RawBlock(text=text, page=current_page)

            if kind:
                block.kind = kind
            else:
                _classify_block(block)

            # Detect is_bold from run formatting (for body paragraphs)
            if block.kind == "body":
                runs = [r for r in para.runs if r.text.strip()]
                if runs:
                    bold_count = sum(1 for r in runs if r.bold)
                    block.is_bold = (bold_count / len(runs)) > 0.5
                    if block.is_bold and len(text) <= MAX_HEADING_CHARS:
                        block.kind = "heading"

            blocks.append(block)

        elif item_type == "table":
            table = item
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ")
                         for cell in row.cells]
                # Deduplicate merged cells (python-docx repeats them)
                seen: list[str] = []
                for c in cells:
                    if not seen or c != seen[-1]:
                        seen.append(c)
                if any(seen):
                    rows.append(" | ".join(seen))
            if rows:
                block = RawBlock(
                    text = "\n".join(rows),
                    page = current_page,
                    kind = "body",   # will be typed as table by layout inference
                )
                blocks.append(block)

    return blocks


# ═══════════════════════════════════════════════════════════════════════════════
# TXT EXTRACTOR
# ═══════════════════════════════════════════════════════════════════════════════

def extract_txt(path: Path) -> list[RawBlock]:
    """
    Extract text blocks from a plain-text file.

    - Form feeds (\\f) mark page boundaries
    - Double newlines separate paragraphs (→ blocks)
    - Single newlines within a paragraph are preserved
    """
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="latin-1")

    blocks: list[RawBlock] = []
    current_page = 1

    # Split on form feeds first to track pages
    pages = raw.split("\f")
    for page_text in pages:
        # Split into paragraphs on double (or more) newlines
        paragraphs = regex.split(r'\n{2,}', page_text)
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            block = RawBlock(text=para, page=current_page)
            _classify_block(block)
            blocks.append(block)
        current_page += 1

    return blocks


# ═══════════════════════════════════════════════════════════════════════════════
# CLAUSE SEGMENTATION
# ═══════════════════════════════════════════════════════════════════════════════

def _is_boundary(block: RawBlock, prev_kind: str) -> bool:
    """
    True if this block should start a new clause.

    Boundaries:
    - A heading block always starts a new clause (if it's not isolated)
    - A colon-heading starts a new clause
    - A table block is always its own clause
    - A body block after a heading = continues the current clause
    """
    return block.kind in ("heading", "heading_colon")


def segment_clauses(blocks: list[RawBlock]) -> list[ClauseCandidate]:
    """
    Convert a flat sequence of classified blocks into clause candidates.

    Rules:
    - Headings start (and optionally open) a new clause
    - A colon-heading is included in the clause it opens (the colon line
      is the clause title; body follows in the same candidate)
    - Tables are always their own clause
    - TOC and skip blocks are discarded
    - Adjacent body/bullet blocks with no heading are bundled together
      up to a natural paragraph break
    """
    active:   ClauseCandidate | None = None
    results:  list[ClauseCandidate]  = []

    # Scan the raw content for inline section breaks (numbered patterns)
    # that appear within body blocks (common in TXT/simple DOCX)
    expanded: list[RawBlock] = []
    for b in blocks:
        if b.kind not in ("body",):
            expanded.append(b)
            continue
        # Try to split body block on numbered headings mid-text
        sub = _split_on_internal_headings(b)
        expanded.extend(sub)

    blocks = expanded

    def _flush():
        nonlocal active
        if active and not active.is_empty:
            results.append(active)
        active = None

    prev_kind = "skip"

    for block in blocks:
        kind = block.kind

        if kind in ("toc", "skip"):
            continue

        # Tables are standalone clauses
        if kind == "body" and _RE_PIPE_ROW.search(block.text):
            # Confirm table-ish content
            pipe_lines = sum(1 for l in block.text.splitlines() if _RE_PIPE_ROW.search(l))
            if pipe_lines >= 2:
                _flush()
                results.append(ClauseCandidate(blocks=[block], page=block.page))
                prev_kind = kind
                continue

        if kind == "heading":
            # Check if next block will attach (we don't look ahead here,
            # so we flush on heading AND let the heading open a new clause)
            _flush()
            active = ClauseCandidate(blocks=[block], page=block.page)

        elif kind == "heading_colon":
            # Colon-heading starts a new clause; body will attach below
            _flush()
            active = ClauseCandidate(blocks=[block], page=block.page)

        else:
            # body / bullet
            if active is None:
                # Orphaned body block: start a new clause with it
                active = ClauseCandidate(blocks=[block], page=block.page)
            else:
                active.blocks.append(block)

        prev_kind = kind

    _flush()
    return results


def _split_on_internal_headings(block: RawBlock) -> list[RawBlock]:
    """
    If a body block contains internal numbered-section headings on their
    own lines, split it into sub-blocks at those boundaries.
    """
    text  = block.text
    lines = text.splitlines(keepends=True)
    if len(lines) < 3:
        return [block]

    # Find lines that look like stand-alone numbered headings
    boundary_indices: list[int] = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if _RE_NUMBERED_HDG.match(stripped) and len(stripped) <= MAX_HEADING_CHARS:
            boundary_indices.append(i)

    if not boundary_indices:
        return [block]

    # Build split positions: start at 0, then each boundary
    splits = sorted(set([0] + boundary_indices))
    sub_blocks: list[RawBlock] = []
    for idx, start in enumerate(splits):
        end = splits[idx + 1] if idx + 1 < len(splits) else len(lines)
        chunk = "".join(lines[start:end]).strip()
        if not chunk:
            continue
        sub = RawBlock(text=chunk, page=block.page)
        # The first line of each chunk determines its kind
        if _RE_NUMBERED_HDG.match(lines[start].strip()):
            sub.kind = "heading"
        else:
            _classify_block(sub)
        sub_blocks.append(sub)

    return sub_blocks if len(sub_blocks) > 1 else [block]


# ═══════════════════════════════════════════════════════════════════════════════
# POST-PROCESSING & OUTPUT BUILDING
# ═══════════════════════════════════════════════════════════════════════════════

def _clean_clause_text(text: str) -> str:
    """Final text cleaning applied to each clause."""
    # Remove lines that are purely artefacts (page numbers, running headers)
    lines = text.splitlines()
    cleaned: list[str] = []
    for line in lines:
        stripped = line.strip()
        if _is_page_artefact(stripped) and stripped:
            continue
        cleaned.append(line)
    text = "\n".join(cleaned)
    # Remove leading/trailing whitespace; collapse 3+ blank lines
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    return text


def _is_heading_only(text: str) -> bool:
    """
    True if the clause text is just a heading title with no body sentences.

    A heading-only clause has a single meaningful line (possibly with a
    trailing colon) and no subsequent body paragraphs.
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if len(lines) != 1:
        return False
    line = lines[0]
    # Must be reasonably short (headings are not paragraphs)
    if len(line) > MAX_HEADING_CHARS:
        return False
    # Must look like a heading: numbered, ALL-CAPS, or ends with colon
    return (
        _is_numbered_heading(line)
        or _is_all_caps_heading(line)
        or _is_colon_heading(line)
    )


def _merge_orphan_headings(candidates: list[ClauseCandidate]) -> list[ClauseCandidate]:
    """
    Merge heading-only candidates into the following candidate.

    When a clause candidate contains only a title line (numbered heading,
    ALL-CAPS heading, or colon-header) with no body text, prepend its
    blocks onto the next candidate so they form a single cohesive clause.

    Example:
        Candidate A: "4. SECURITY INCIDENT RESPONSE"  ← heading only
        Candidate B: "Security incident...obligations:\n• item1\n• item2"
        → merged: Candidate with A.blocks + B.blocks, page = A.page
    """
    if not candidates:
        return candidates

    merged: list[ClauseCandidate] = []
    i = 0
    while i < len(candidates):
        cand = candidates[i]
        text = _normalise(_clean_clause_text(cand.text))

        if _is_heading_only(text) and i + 1 < len(candidates):
            # Absorb this heading into the next candidate
            next_cand = candidates[i + 1]
            combined  = ClauseCandidate(
                blocks = cand.blocks + next_cand.blocks,
                page   = cand.page,
            )
            merged.append(combined)
            i += 2          # skip both
        else:
            merged.append(cand)
            i += 1

    return merged


def build_output(candidates: list[ClauseCandidate]) -> list[dict]:
    """
    Convert clause candidates to the final output list.

    - Merges orphan heading-only clauses with the following clause
    - Assigns CL-NNN IDs (1-indexed, zero-padded to 3 digits)
    - Infers layout_type
    - Discards clauses shorter than MIN_CLAUSE_CHARS
    """
    # Post-processing: merge standalone headings into subsequent body
    candidates = _merge_orphan_headings(candidates)

    output: list[dict] = []
    seq = 1

    for candidate in candidates:
        text = _normalise(_clean_clause_text(candidate.text))
        if len(text.strip()) < MIN_CLAUSE_CHARS:
            continue

        layout = _infer_layout(text)

        clause: dict = {
            "clause_id":   f"CL-{seq:03d}",
            "page":        candidate.page,
            "layout_type": layout,
            "text":        text,
        }
        output.append(clause)
        seq += 1

    return output


# ═══════════════════════════════════════════════════════════════════════════════
# LLM-ASSISTED SEGMENTATION
# ═══════════════════════════════════════════════════════════════════════════════

# Maximum characters of raw text to send in a single LLM segmentation call.
# GPT-4o / GPT-4.5 support 128K tokens; ~90K chars is a safe limit (~60K tokens).
_LLM_SEG_CHUNK_CHARS = 90_000
_LLM_SEG_OVERLAP_CHARS = 4_000   # overlap between chunks to avoid splitting a §


def _llm_segment_clauses(raw_text: str, llm_provider: object) -> list[dict] | None:
    """
    Use the configured LLM to identify all InfoSec/compliance-relevant passages
    in *raw_text*.  Returns a list of clause dicts or None on failure.

    Each returned dict has: clause_id, page, layout_type, text,
    category, relevance_reason.
    """
    try:
        from llm.prompts import (
            SEGMENTATION_SYSTEM_PROMPT,
            SEGMENTATION_OUTPUT_SCHEMA,
            PROMPT_VERSION_SEGMENTATION,
            build_segmentation_user_message,
        )
    except ImportError:
        return None

    # ── Split into chunks if the text is very long ────────────────────────────
    chunks: list[str] = []
    if len(raw_text) <= _LLM_SEG_CHUNK_CHARS:
        chunks = [raw_text]
    else:
        start = 0
        while start < len(raw_text):
            end = start + _LLM_SEG_CHUNK_CHARS
            chunks.append(raw_text[start:end])
            start = end - _LLM_SEG_OVERLAP_CHARS

    all_segments: list[dict] = []
    for chunk in chunks:
        user_msg = build_segmentation_user_message(chunk)
        try:
            response = llm_provider.complete_structured(
                system_prompt  = SEGMENTATION_SYSTEM_PROMPT,
                user_message   = user_msg,
                json_schema    = SEGMENTATION_OUTPUT_SCHEMA,
                prompt_version = PROMPT_VERSION_SEGMENTATION,
                max_tokens     = 4096,
            )
        except Exception:
            return None  # any LLM error → fall back to rule-based

        if response is None:
            return None

        segments = (response.content or {}).get("segments", [])
        if not isinstance(segments, list):
            return None
        all_segments.extend(segments)

    if not all_segments:
        return None

    # ── Build clause dicts compatible with stage4_clauses.json schema ─────────
    clauses: list[dict] = []
    seen_texts: set[str] = set()
    seq = 1
    for seg in all_segments:
        text = _normalise(str(seg.get("text", "")).strip())
        if len(text) < MIN_CLAUSE_CHARS:
            continue
        # De-duplicate overlapping chunks that may produce the same passage twice
        key = text[:120]
        if key in seen_texts:
            continue
        seen_texts.add(key)

        category = seg.get("category") or "other_relevant"
        clauses.append({
            "clause_id":        f"CL-{seq:03d}",
            "page":             int(seg.get("page_hint") or 1),
            "layout_type":      category,          # reuse layout_type slot for category
            "text":             text,
            "category":         category,
            "relevance_reason": str(seg.get("relevance_reason") or ""),
        })
        seq += 1

    return clauses if clauses else None


# ═══════════════════════════════════════════════════════════════════════════════
# PUBLIC API
# ═══════════════════════════════════════════════════════════════════════════════

SUPPORTED_FORMATS = {".pdf", ".docx", ".txt"}


def ingest(contract_path: str | Path, llm_provider: object = None) -> list[dict]:
    """
    Main entry point.  Convert a contract file to the stage4_clauses.json list.

    When *llm_provider* is supplied the LLM performs intelligent segmentation:
    it reads the full contract text and returns all InfoSec/compliance-relevant
    passages regardless of document structure (§, Article, numbered sections…).
    Falls back to rule-based heading segmentation when the LLM is unavailable.

    Parameters
    ----------
    contract_path : str | Path
        Path to the contract file (.pdf, .docx, or .txt)
    llm_provider : BaseLLMProvider | None
        Configured LLM provider for intelligent segmentation (optional).

    Returns
    -------
    list[dict]
        List of clause dicts compatible with stage4_clauses.json schema.

    Raises
    ------
    FileNotFoundError
        If the input file does not exist.
    ValueError
        If the file extension is not supported.
    """
    path = Path(contract_path)
    if not path.exists():
        raise FileNotFoundError(f"Contract file not found: {path}")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_FORMATS:
        raise ValueError(
            f"Unsupported format '{suffix}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_FORMATS))}"
        )

    # ── 1. Extract raw blocks ──────────────────────────────────────────────────
    if suffix == ".pdf":
        blocks = extract_pdf(path)
    elif suffix == ".docx":
        blocks = extract_docx(path)
    else:  # .txt
        blocks = extract_txt(path)

    # ── 2. Normalise each block ────────────────────────────────────────────────
    for b in blocks:
        b.text = _normalise(b.text)

    # ── 3a. LLM-assisted segmentation (preferred) ─────────────────────────────
    if llm_provider is not None:
        raw_text = "\n\n".join(b.text for b in blocks if b.text.strip())
        llm_clauses = _llm_segment_clauses(raw_text, llm_provider)
        if llm_clauses:
            return llm_clauses
        # LLM returned nothing useful → fall through to rule-based

    # ── 3b. Rule-based segmentation (fallback) ────────────────────────────────
    candidates = segment_clauses(blocks)

    # ── 4. Build and return output ─────────────────────────────────────────────
    return build_output(candidates)


# ═══════════════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        prog        = "stage16_contract_ingestion",
        description = "Stage 16 — Convert raw contract to stage4_clauses.json.",
        formatter_class = argparse.RawDescriptionHelpFormatter,
        epilog = (
            "Supported formats: PDF, DOCX, TXT\n\n"
            "Examples:\n"
            "  %(prog)s --contract contract.pdf\n"
            "  %(prog)s --contract agreement.docx --output clauses.json\n"
            "  %(prog)s --contract terms.txt --quiet\n"
        ),
    )
    p.add_argument(
        "--contract", "-c",
        required = True,
        metavar  = "FILE",
        help     = "Path to the contract file (.pdf, .docx, or .txt).",
    )
    p.add_argument(
        "--output", "-o",
        default  = "stage4_clauses.json",
        metavar  = "FILE",
        help     = "Output JSON file (default: stage4_clauses.json).",
    )
    p.add_argument(
        "--quiet", "-q",
        action  = "store_true",
        help    = "Suppress summary output.",
    )
    p.add_argument(
        "--pretty",
        action  = "store_true",
        default = True,
        help    = "Pretty-print JSON output (default: on).",
    )
    return p.parse_args()


def _print_summary(clauses: list[dict], output_path: str, source: str) -> None:
    layout_counts: dict[str, int] = {}
    for c in clauses:
        lt = c.get("layout_type", "unknown")
        layout_counts[lt] = layout_counts.get(lt, 0) + 1

    pages = [c["page"] for c in clauses if c.get("page")]
    min_pg = min(pages) if pages else "—"
    max_pg = max(pages) if pages else "—"

    print(f"\n{'='*64}")
    print(f"  STAGE 16 — CONTRACT INGESTION")
    print(f"{'='*64}")
    print(f"  Source  : {source}")
    print(f"  Output  : {output_path}")
    print(f"  Clauses : {len(clauses)}")
    print(f"  Pages   : {min_pg} – {max_pg}")
    print(f"  Layout breakdown:")
    for lt, n in sorted(layout_counts.items()):
        print(f"    {lt:<16} {n:>3}")
    print(f"{'='*64}\n")

    if clauses:
        print("  First 3 clauses:")
        for c in clauses[:3]:
            preview = c["text"][:80].replace("\n", " ")
            print(f"    {c['clause_id']} (p.{c['page']} {c['layout_type']}): {preview}…")
        print()


def main() -> None:
    args = _parse_args()

    try:
        clauses = ingest(args.contract)
    except FileNotFoundError as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(2)
    except ValueError as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(2)
    except Exception as e:
        print(f"[ERROR] Unexpected failure during ingestion: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        sys.exit(2)

    indent = 2 if args.pretty else None
    with open(args.output, "w", encoding="utf-8") as f:
        json.dump(clauses, f, indent=indent, ensure_ascii=False)

    if not args.quiet:
        _print_summary(clauses, args.output, args.contract)

    sys.exit(0)


if __name__ == "__main__":
    main()
